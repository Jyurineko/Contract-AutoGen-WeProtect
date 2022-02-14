from tkinter import *
from docx import Document
import os


# def print_hi(name):
#     # 在下面的代码行中使用断点来调试脚本。
#     print(f'Hi, {name}')  # 按 Ctrl+F8 切换断点。
#
#
# # 按间距中的绿色按钮以运行脚本。
# if __name__ == '__main__':
#     print_hi('PyCharm')

# https://www.codegrepper.com/code-examples/python/python+tkinter+text+get

def main():
    # exec 动态函数调用 勿删
    def focus_next_window(event):
        event.widget.tk_focusNext().focus()
        return "break"

    def shuttle_text(shuttle):
        t = ''
        for i in shuttle:
            t += i.text
        return t

    def docx_replace(doc, data):
        for key in data:
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if key in cell.text:
                            cell.text = cell.text.replace(key, data[key])

            for p in doc.paragraphs:

                begin = 0
                for end in range(len(p.runs)):

                    shuttle = p.runs[begin:end + 1]

                    full_text = shuttle_text(shuttle)
                    if key in full_text:
                        # print('Replace：', key, '->', data[key])
                        # print([i.text for i in shuttle])

                        # find the begin
                        index = full_text.index(key)
                        # print('full_text length', len(full_text), 'index:', index)
                        while index >= len(p.runs[begin].text):
                            index -= len(p.runs[begin].text)
                            begin += 1

                        shuttle = p.runs[begin:end + 1]

                        # do replace
                        # print('before replace', [i.text for i in shuttle])
                        if key in shuttle[0].text:
                            shuttle[0].text = shuttle[0].text.replace(key, data[key])
                        else:
                            replace_begin_index = shuttle_text(shuttle).index(key)
                            replace_end_index = replace_begin_index + len(key)
                            replace_end_index_in_last_run = replace_end_index - len(shuttle_text(shuttle[:-1]))
                            shuttle[0].text = shuttle[0].text[:replace_begin_index] + data[key]

                            # clear middle runs
                            for i in shuttle[1:-1]:
                                i.text = ''

                            # keep last run
                            shuttle[-1].text = shuttle[-1].text[replace_end_index_in_last_run:]

                        # print('after replace', [i.text for i in shuttle])

                        # set begin to next
                        begin = end

    def word_gen():
        new_name = textbox_name.get("1.0", "end-1c")
        user_id = textbox_uid.get("1.0", "end-1c")
        new_date = textbox_date.get("1.0", "end-1c")
        new_addr = textbox_addr.get("1.0", "end-1c")
        new_currency = textbox_curr.get("1.0", "end-1c")
        new_city = textbox_city.get("1.0", "end-1c")
        new_laws = textbox_laws.get("1.0", "end-1c")
        new_effect = textbox_effe.get("1.0", "end-1c")
        new_notice = textbox_noti.get("1.0", "end-1c")
        new_compen = textbox_comp.get("1.0", "end-1c")

        doc_name = user_id
        password = user_id

        print(new_name, new_date, new_addr, new_city, new_laws, new_currency, new_effect, new_notice,
              new_compen + '获取正确!!')

        docx_replace(doc, dict({old_date: new_date, old_name: new_name, old_addr: new_addr, old_city: new_city,
                                old_laws: new_laws, old_currency: new_currency, old_effect: new_effect,
                                old_notice: new_notice, old_compen: new_compen}))

        filename = 'UserID_' + doc_name + 'Independent_Contractor_Agreement.docx'
        filepath = './Generated-Contracts/'
        doc.save(filepath + filename)

        #  加密word，密码是user ID

        #  打开word
        os.system('start' + filepath + filename)

    def menu_New():
        textbox_name.delete("1.0", "end-1c")
        textbox_uid.delete("1.0", "end-1c")
        textbox_date.delete("1.0", "end-1c")
        textbox_addr.delete("1.0", "end-1c")
        textbox_curr.delete("1.0", "end-1c")
        textbox_city.delete("1.0", "end-1c")
        textbox_laws.delete("1.0", "end-1c")
        textbox_effe.delete("1.0", "end-1c")
        textbox_noti.delete("1.0", "end-1c")
        textbox_comp.delete("1.0", "end-1c")

    def menu_consume_fraud():
        consume_fraud_label = ['User(用户):', 'UID(用户ID)', 'Date(日期):', 'Address(地址):', 'Currency(货币):',
                               'City,Zip-code(城市,邮编):', 'Laws of the State(管辖权):']
        i = 0
        for con in consume_fraud_label:
            Label(text=con).grid(column=0, row=i)
            i += 1

    win = Tk()
    win.title('WeProtect内部合同生成系统/Contract Gen System')
    win.geometry("800x300")

    #  创建菜单栏 在窗口的上方
    menubar = Menu(win)

    #  创建新菜单栏
    menu_file = Menu(menubar, tearoff=0)
    menu_personal_law = Menu(menubar, tearoff=0)
    menu_business_law = Menu(menubar, tearoff=0)

    # 命名空菜单，放进菜单栏中
    menubar.add_cascade(label='File', menu=menu_file)
    menubar.add_cascade(label='Personal_Law', menu=menu_personal_law)
    menubar.add_cascade(label='Business_Law', menu=menu_business_law)

    # 在File菜单中，添加New\open\exit等小菜单，每一个小菜单对应命令操作
    menu_file.add_command(label='New', command=menu_New)
    menu_file.add_command(label='Quit', command=win.quit)

    # 个人法律菜单栏的业务作为小菜单，每个小菜单对应一个新的windows，里面的要填的内容需要随时更改
    personal_law_product = ['消费纠纷', '民事纠纷', '移民签证', '劳务纠纷', '交通罚单', '刑事案件', '法律文件撰写', '小额法庭',
                            '合同审核']
    for p in personal_law_product:
        menu_personal_law.add_command(label=p, command=menu_consume_fraud)  # todo 每个个人业务要有一个新command 改变填写内容

    # 商业法律菜单栏的业务作为小菜单，每个小菜单对应一个新的windows，里面的要填的内容需要随时更改
    business_law_product = ['知识产权保护', '海关事务', '产品合规', '商业合同', '跨境专项服务']
    for bus in business_law_product:
        menu_business_law.add_command(label=bus, command=menu_consume_fraud)  # todo 每个商业业务要有一个新command 改变填写内容

    # 配置 显示菜单栏menubar
    win.config(menu=menubar)

    #  grid label
    label_name = ['User(用户):', 'UID(用户ID)', 'Date(日期):', 'Address(地址):', 'Currency(货币):', 'City,Zip-code(城市,邮编):',
                  'Laws of the State(管辖权):', 'Effectiveness(时效(月份)):', 'Days of Notice Ahead(提前几日通知):',
                  'Compensation(酬劳条款):']

    textbox_var = ['name', 'uid', 'date', 'addr', 'curr', 'city', 'laws', 'effe', 'noti', 'comp']

    # 生成label名字
    i = 0
    for a in label_name:
        Label(text=a).grid(column=0, row=i)
        i = i + 1

    # 生成对应的textbox
    textbox_name = Text(win, height=1)
    textbox_uid = Text(win, height=1)
    textbox_date = Text(win, height=1)
    textbox_addr = Text(win, height=1)
    textbox_curr = Text(win, height=1)
    textbox_city = Text(win, height=1)
    textbox_laws = Text(win, height=1)
    textbox_effe = Text(win, height=1)
    textbox_noti = Text(win, height=1)
    textbox_comp = Text(win, height=3)

    x = 0
    for b in textbox_var:
        exec('textbox_{}.grid(column=1, row={})'.format(b, x))
        exec('textbox_{}.bind("<Tab>", focus_next_window)'.format(b))
        x = x + 1

    old_date = 'AIdate'
    old_name = 'AIname'
    old_addr = 'AIroadaddress'
    old_city = 'AIcityandcode'
    old_laws = 'AIlawsstate'
    old_currency = 'AIcurrency'
    old_effect = 'AIeffect'
    old_notice = 'AInotice'
    old_compen = 'AIcompen'

    doc = Document('./basefile/Contract.docx')

    #  word生成 按钮
    # Button(win, text='合同生成', command=word_gen).place(x=90, y=380)
    Button(win, text='合同生成(Gen Contract)', command=word_gen).grid(column=1, row=14)
    win.iconbitmap('./logo.ico')
    win.mainloop()  # 进入消息循环


if __name__ == '__main__':
    main()
